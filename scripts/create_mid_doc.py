#!/usr/bin/env python3
"""
create_mid_doc.py
Generate the 2-page written mid-project update as a Word (.docx) file.

Run:    python3 create_mid_doc.py
Output: results/mid_project_update.docx
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


doc = Document()

# Tight 2-page layout
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ---------- Title ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Mid-Project Update  —  Anti-Sloshing Controller for a Food-Serving Robot")
r.bold = True
r.font.size = Pt(15)

byline = doc.add_paragraph()
byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = byline.add_run("Hyundae Cha   |   MECE 6397   |   2026-04-15")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def H(text, level=2):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12 if level == 2 else 11)
    r.font.color.rgb = RGBColor(0x1A, 0x50, 0xAA)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def P(runs):
    """runs = list of (text, bold, italic) tuples, or a single string."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if isinstance(runs, str):
        p.add_run(runs)
        return p
    for text, bold, italic in runs:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    return p


def bullet(runs, indent=0.20):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(indent)
    if isinstance(runs, str):
        p.add_run(runs)
        return p
    for text, bold, italic in runs:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    return p


# ============================================================
# 1. Objective and Motivation
# ============================================================
H("1.  Objective and Motivation")
P("Serving robots carry open containers of liquid across restaurants. Even "
  "small accelerations excite sloshing, which causes spills and forces "
  "operators to drive the robot slowly — eroding the productivity case for "
  "automation. The goal of this project is to reproduce, analyze, and extend "
  "the Self-Balancing Slosh-Free Controller (SBSFC) of Choi et al. (2024) so "
  "that a two-wheeled robot can transport liquid at useful speeds without "
  "sloshing.")

P([("The core physical insight is that sloshing is driven by ", False, False),
   ("acceleration, not velocity", True, True),
   (". A fast but smooth ramp spills nothing; a slow but jerky motion spills "
    "everything. Any controller that only smooths the velocity command "
    "(e.g. a low-pass filter, LPF) leaves the sloshing-band content of the "
    "acceleration untouched. SBSFC targets that band directly.", False, False)])

# ============================================================
# 2. Progress to Date
# ============================================================
H("2.  Progress to Date")
P("A full MATLAB simulation environment has been built and validated against "
  "the paper's Scenario 1 (sudden start / sudden stop). The following "
  "components are complete:")

bullet([("Plant model. ", True, False),
        ("Linearized two-wheel self-balancing robot with body tilt ψ, plus a "
         "pendulum-equivalent sloshing model with natural frequency "
         "ω_f = √(g/l) ≈ 8.09 rad/s (1.29 Hz) and light viscous damping.",
         False, False)])

bullet([("Baseline controller (LPF). ", True, False),
        ("1st-order Butterworth low-pass filter on the velocity reference — "
         "the straw-man the paper compares against.", False, False)])

bullet([("SBSFC controller. ", True, False),
        ("Three blocks reproduced faithfully: (i) input-shaping filter F_c(s) "
         "whose zeros cancel the sloshing poles; (ii) LQT state-feedback gain "
         "K from the Riccati equation (state [x, ψ, ẋ, ψ̇], "
         "Q = diag(500, 50, 1, 1), R = 0.1); (iii) DOB + compensator that "
         "estimates the residual tilt disturbance from ψ̈ and adds a "
         "sign(ψ̇_filt) sliding-mode correction.", False, False)])

bullet([("Test harness. ", True, False),
        ("Scripted Scenario 1 (step to 0.6 m/s, hold, step to 0, idle, step "
         "to 0.3 m/s, stop), deterministic disturbance, Euler integration at "
         "dt = 1 ms, 30 s horizon.", False, False)])

bullet([("Analysis outputs. ", True, False),
        ("Table 3 reproduction, the 4-panel velocity-chain figure, and an "
         "interactive 2-D animation with pause / resume / speed slider.",
         False, False)])

# ============================================================
# 3. Preliminary Results
# ============================================================
H("3.  Preliminary Results  (Scenario 1)")

table = doc.add_table(rows=4, cols=4)
table.style = "Light Grid Accent 1"
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells = table.rows[0].cells
headers = ["Metric", "With SBSFC", "Without SBSFC", "Reduction"]
for i, h in enumerate(headers):
    hdr_cells[i].text = ""
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold = True; r.font.size = Pt(10)
    set_cell_shading(hdr_cells[i], "1A50AA")
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows = [
    ("Mean  |θ|   [deg]",          "0.317",  "5.066",  "93.7 %"),
    ("Variance  θ   [deg²]",       "0.193",  "35.413", "99.5 %"),
    ("Peak body acceleration  [m/s²]", "0.20", "1.14",  "5.7× smaller"),
]
for ri, row in enumerate(rows, start=1):
    for ci, val in enumerate(row):
        cell = table.rows[ri].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val); r.font.size = Pt(10)
        if ci == 3:
            r.bold = True
            r.font.color.rgb = RGBColor(0x1E, 0x6E, 0x2E)

P([("The sloshing-angle reductions reproduce the paper's figures almost "
    "exactly for this scenario. More informative is the mechanism exposed by "
    "the acceleration / FFT analysis: the two velocity traces ẋ(t) look "
    "nearly identical, yet the LPF produces sharp acceleration spikes at "
    "every step while SBSFC spreads and ", False, False),
   ("notches", True, False),
   (" the acceleration spectrum precisely at the sloshing resonance "
    "ω_f/2π = 1.29 Hz. This is the clearest single piece of evidence that "
    "SBSFC is doing something fundamentally different from “just a smoother "
    "filter”.", False, False)])

# ============================================================
# 4. Roadblocks
# ============================================================
H("4.  Roadblocks and Limitations")
P("Reproducing the paper surfaced four non-trivial limitations that must be "
  "addressed before any hardware step:")

bullet([("Hand-tuned gains. ", True, False),
        ("K, η, Q, R, and the compensator thresholds are tuned for one "
         "plant. Any change in liquid, container, or robot requires "
         "re-tuning by hand.", False, False)])
bullet([("Known ω_f assumption. ", True, False),
        ("The input shaper needs ω_f exactly. In reality ω_f drifts with "
         "fill level, temperature, and container shape, and a fixed notch "
         "mis-tunes within seconds.", False, False)])
bullet([("Narrow test set. ", True, False),
        ("Only five synthetic 1-D scenarios have been tested. Real delivery "
         "paths are 2-D and include curves, aisles, and turns around tables.",
         False, False)])
bullet([("No hardware validation. ", True, False),
        ("All results are in simulation; wheel stick-slip, motor delay, IMU "
         "noise, and true liquid physics are not yet in the loop.",
         False, False)])

# ============================================================
# 5. Plan
# ============================================================
H("5.  Plan for the Remainder of the Semester")
P("The remaining work targets the first two limitations, which are the ones "
  "that block any deployment:")

bullet([("RL controller (primary extension). ", True, False),
        ("Train a model-free policy (PPO or SAC) on the same MATLAB plant, "
         "using state [x, ẋ, ψ, ψ̇, θ, θ̇] and reward "
         "−(α·θ² + β·tracking_error²). The goal is a policy that (i) needs "
         "no hand-tuning and (ii) adapts across a randomized ω_f and "
         "fill-level distribution via domain randomization — directly "
         "attacking limitations 1 and 2. SBSFC serves as the quantitative "
         "baseline.", False, False)])
bullet([("Expanded scenarios. ", True, False),
        ("Add curved paths and 90° turns around tables, vary fill level and "
         "viscosity, and include adversarial disturbances. This upgrades the "
         "evaluation from “reproduces the paper” to “predicts real serving "
         "workloads”.", False, False)])
bullet([("Stretch: 2-D motion. ", True, False),
        ("Extend the plant and controller to planar motion so the same RL "
         "policy can be evaluated on restaurant-like trajectories rather "
         "than 1-D profiles.", False, False)])

P([("Timeline. ", True, False),
   ("RL training pipeline set up by late April; first PPO vs SBSFC comparison "
    "on Scenario 1 by early May; randomized-ω_f evaluation and final report "
    "by the project due date.", False, False)])

# ============================================================
# 6. Risk Assessment
# ============================================================
H("6.  Risk Assessment")
P("The RL sample-efficiency on a 6-state plant with a second-order sloshing "
  "sub-system is the main technical risk. If PPO fails to converge within a "
  "reasonable simulation budget, the fallback is adaptive SBSFC — a "
  "recursive-least-squares estimator for ω_f feeding the existing shaper "
  "F_c(s) — which addresses limitation 2 without requiring a new controller "
  "class.")

# ============================================================
# References
# ============================================================
H("References")

def ref(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text); r.font.size = Pt(9.5)

ref("[1]  J. Choi et al., “Self-Balancing Slosh-Free Control of a Two-Wheeled "
    "Food-Serving Robot,” Mechatronics, 2024.")
ref("[2]  B. Siciliano and O. Khatib (eds.), Springer Handbook of Robotics, "
    "2nd ed., Ch. 25 (Wheeled Robots) and Ch. 8 (Motion Control).")
ref("[3]  J. Schulman et al., “Proximal Policy Optimization Algorithms,” "
    "arXiv:1707.06347, 2017.")
ref("[4]  T. Haarnoja et al., “Soft Actor-Critic,” ICML 2018.")

# ============================================================
# Save
# ============================================================
os.makedirs("results", exist_ok=True)
out = "results/mid_project_update.docx"
doc.save(out)
print(f"Saved: {out}")
